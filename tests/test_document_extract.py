"""Tests for document extraction (PDF, DOCX, Markdown)."""

import io
from unittest.mock import patch

import pytest

from aion.tools.document_extract import (
    ExtractedDocument,
    extract_docx,
    extract_markdown,
    extract_pdf,
)


# ---------------------------------------------------------------------------
# Helpers to create minimal test files in memory
# ---------------------------------------------------------------------------


def _make_pdf_bytes(text: str = "Hello from test PDF") -> bytes:
    """Create a minimal valid PDF with text content using PyMuPDF."""
    import fitz

    doc = fitz.open()
    page = doc.new_page()
    page.insert_text((72, 72), text)
    data = doc.tobytes()
    doc.close()
    return data


def _make_docx_bytes(
    paragraphs: list[str] | None = None,
    tables: list[list[list[str]]] | None = None,
) -> bytes:
    """Create a minimal valid DOCX with paragraphs and optional tables."""
    from docx import Document

    doc = Document()
    for p in (paragraphs or ["Hello from test DOCX"]):
        doc.add_paragraph(p)
    if tables:
        for table_data in tables:
            rows, cols = len(table_data), len(table_data[0])
            tbl = doc.add_table(rows=rows, cols=cols)
            for i, row_data in enumerate(table_data):
                for j, cell_text in enumerate(row_data):
                    tbl.cell(i, j).text = cell_text
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


# ===========================================================================
# PDF extraction
# ===========================================================================


class TestExtractPdf:

    def test_valid_pdf(self):
        data = _make_pdf_bytes("Test document content here")
        result = extract_pdf(data)
        assert isinstance(result, ExtractedDocument)
        assert "Test document content here" in result.content
        assert result.file_type == "pdf"
        assert result.page_count >= 1
        assert result.word_count > 0
        assert result.char_count > 0

    def test_empty_bytes_raises(self):
        with pytest.raises(ValueError, match="Empty PDF"):
            extract_pdf(b"")

    def test_corrupt_bytes_raises(self):
        with pytest.raises(ValueError):
            extract_pdf(b"not a pdf at all")

    def test_multi_page_pdf(self):
        import fitz

        doc = fitz.open()
        for i in range(3):
            page = doc.new_page()
            page.insert_text((72, 72), f"Page {i + 1} content")
        data = doc.tobytes()
        doc.close()

        result = extract_pdf(data)
        assert result.page_count == 3
        assert "Page 1" in result.content
        assert "Page 3" in result.content

    def test_title_from_first_line(self):
        data = _make_pdf_bytes("Architecture Decision Record\nSome details")
        result = extract_pdf(data)
        assert "Architecture Decision Record" in result.title


# ===========================================================================
# DOCX extraction
# ===========================================================================


class TestExtractDocx:

    def test_valid_docx(self):
        data = _make_docx_bytes(["First paragraph", "Second paragraph"])
        result = extract_docx(data)
        assert isinstance(result, ExtractedDocument)
        assert "First paragraph" in result.content
        assert "Second paragraph" in result.content
        assert result.file_type == "docx"
        assert result.word_count > 0

    def test_empty_bytes_raises(self):
        with pytest.raises(ValueError, match="Empty DOCX"):
            extract_docx(b"")

    def test_corrupt_bytes_raises(self):
        with pytest.raises(ValueError, match="Failed to read DOCX"):
            extract_docx(b"not a docx file")

    def test_with_tables(self):
        table = [["Header A", "Header B"], ["Cell 1", "Cell 2"]]
        data = _make_docx_bytes(
            paragraphs=["Intro text"],
            tables=[table],
        )
        result = extract_docx(data)
        assert "Header A" in result.content
        assert "Cell 1" in result.content

    def test_empty_docx_raises(self):
        """DOCX with no text content should raise."""
        from docx import Document

        doc = Document()
        buf = io.BytesIO()
        doc.save(buf)
        with pytest.raises(ValueError, match="empty"):
            extract_docx(buf.getvalue())


# ===========================================================================
# Markdown extraction
# ===========================================================================


class TestExtractMarkdown:

    def test_with_frontmatter(self):
        md = """---
title: Test Principle
status: proposed
---

# Heading

Some content here.
"""
        result = extract_markdown(md)
        assert result.title == "Test Principle"
        assert "Some content here" in result.content
        assert result.file_type == "md"

    def test_without_frontmatter(self):
        md = """# My Document

This is the body text.
"""
        result = extract_markdown(md)
        assert result.title == "My Document"
        assert "body text" in result.content

    def test_plain_text(self):
        md = "Just plain text with no headings or frontmatter."
        result = extract_markdown(md)
        assert "plain text" in result.content
        assert result.word_count > 0

    def test_empty_raises(self):
        with pytest.raises(ValueError, match="Empty"):
            extract_markdown("")

    def test_whitespace_only_raises(self):
        with pytest.raises(ValueError, match="Empty"):
            extract_markdown("   \n\n  ")

    def test_frontmatter_name_field(self):
        md = """---
name: My Named Doc
---

Content.
"""
        result = extract_markdown(md)
        assert result.title == "My Named Doc"


# ===========================================================================
# Upload endpoint integration (if chat_ui is importable)
# ===========================================================================


class TestUploadEndpoint:
    """Integration tests for /api/chat/upload with document types."""

    @pytest.fixture
    def client(self, tmp_path):
        """Create a test client with mocked lifespan and temp DB."""
        from contextlib import asynccontextmanager

        import aion.chat_ui as chat_ui_mod

        @asynccontextmanager
        async def _mock_lifespan(a):
            yield

        original_lifespan = chat_ui_mod.app.router.lifespan_context
        original_db = chat_ui_mod._db_path
        test_db = tmp_path / "test_upload.db"

        chat_ui_mod._db_path = test_db
        chat_ui_mod.app.router.lifespan_context = _mock_lifespan
        chat_ui_mod.init_db()

        try:
            from starlette.testclient import TestClient

            with TestClient(chat_ui_mod.app) as c:
                yield c
        finally:
            chat_ui_mod.app.router.lifespan_context = original_lifespan
            chat_ui_mod._db_path = original_db

    def test_upload_pdf(self, client):
        data = _make_pdf_bytes("Test PDF content for upload")
        resp = client.post(
            "/api/chat/upload",
            files={"file": ("test.pdf", data, "application/pdf")},
        )
        assert resp.status_code == 200
        body = resp.json()
        assert body["filename"] == "test.pdf"
        assert body["file_type"] == "pdf"
        assert body["page_count"] >= 1
        assert body["word_count"] > 0
        assert "conversation_id" in body
        assert "artifact_id" in body

    def test_upload_docx(self, client):
        data = _make_docx_bytes(["DOCX upload test paragraph"])
        resp = client.post(
            "/api/chat/upload",
            files={"file": ("test.docx", data, "application/vnd.openxmlformats-officedocument.wordprocessingml.document")},
        )
        assert resp.status_code == 200
        body = resp.json()
        assert body["filename"] == "test.docx"
        assert body["file_type"] == "docx"

    def test_upload_markdown(self, client):
        md = "# Test Upload\n\nSome markdown content."
        resp = client.post(
            "/api/chat/upload",
            files={"file": ("test.md", md.encode(), "text/markdown")},
        )
        assert resp.status_code == 200
        body = resp.json()
        assert body["filename"] == "test.md"
        assert body["file_type"] == "md"

    def test_upload_unsupported_type(self, client):
        resp = client.post(
            "/api/chat/upload",
            files={"file": ("test.exe", b"binary", "application/octet-stream")},
        )
        assert resp.status_code == 400
        assert "Unsupported" in resp.json()["detail"]

    def test_upload_too_large(self, client):
        large_data = _make_pdf_bytes("x" * 100)
        # Patch threshold to 100 bytes so the test PDF exceeds it
        with patch(
            "aion.chat_ui.get_thresholds_value",
            return_value={"max_file_bytes": 100, "max_text_chars": 500_000},
        ):
            resp = client.post(
                "/api/chat/upload",
                files={"file": ("big.pdf", large_data, "application/pdf")},
            )
        assert resp.status_code == 400
        assert "limit" in resp.json()["detail"].lower()
