"""Extract text from uploaded documents (PDF, DOCX, Markdown).

In-memory API: accepts bytes (PDF/DOCX) or str (Markdown).
No file-path dependencies, no Weaviate concerns.
"""

from __future__ import annotations

import io
import re
from dataclasses import dataclass

import structlog

logger = structlog.get_logger(__name__)


@dataclass
class ExtractedDocument:
    """Result of text extraction from an uploaded document."""

    title: str
    content: str
    file_type: str  # "pdf", "docx", "md"
    page_count: int  # 0 for non-PDF
    word_count: int
    char_count: int


def extract_pdf(data: bytes) -> ExtractedDocument:
    """Extract text from PDF bytes using PyMuPDF, fallback to pypdf."""
    if not data:
        raise ValueError("Empty PDF data")

    text = _extract_pdf_pymupdf(data)
    page_count = _count_pdf_pages(data)

    if not text or not text.strip():
        # Fallback to pypdf
        text, page_count = _extract_pdf_pypdf(data)

    if not text or not text.strip():
        raise ValueError(
            "Document appears to be empty or contains only images. "
            "Text extraction found no readable content."
        )

    title = _title_from_text(text, fallback="Uploaded PDF")
    return ExtractedDocument(
        title=title,
        content=text.strip(),
        file_type="pdf",
        page_count=page_count,
        word_count=len(text.split()),
        char_count=len(text),
    )


def _extract_pdf_pymupdf(data: bytes) -> str:
    """Extract text using PyMuPDF (fitz) in-memory API."""
    try:
        import fitz

        doc = fitz.open(stream=data, filetype="pdf")
        pages = []
        for page in doc:
            pages.append(page.get_text())
        doc.close()
        return "\n\n".join(pages)
    except Exception as exc:
        logger.warning("pymupdf_extraction_failed", error=str(exc))
        return ""


def _count_pdf_pages(data: bytes) -> int:
    try:
        import fitz

        doc = fitz.open(stream=data, filetype="pdf")
        count = len(doc)
        doc.close()
        return count
    except Exception:
        return 0


def _extract_pdf_pypdf(data: bytes) -> tuple[str, int]:
    """Fallback PDF extraction using pypdf."""
    try:
        from pypdf import PdfReader

        reader = PdfReader(io.BytesIO(data))
        pages = []
        for page in reader.pages:
            text = page.extract_text()
            if text:
                pages.append(text)
        return "\n\n".join(pages), len(reader.pages)
    except Exception as exc:
        logger.warning("pypdf_extraction_failed", error=str(exc))
        raise ValueError(f"Failed to read PDF: {exc}") from exc


def extract_docx(data: bytes) -> ExtractedDocument:
    """Extract text from DOCX bytes using python-docx."""
    if not data:
        raise ValueError("Empty DOCX data")

    try:
        from docx import Document
    except ImportError as exc:
        raise ValueError("python-docx is not installed") from exc

    try:
        doc = Document(io.BytesIO(data))
    except Exception as exc:
        raise ValueError(f"Failed to read DOCX: {exc}") from exc

    parts: list[str] = []

    # Extract paragraphs
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            parts.append(text)

    # Extract table content
    for table in doc.tables:
        rows = []
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            rows.append(" | ".join(cells))
        if rows:
            parts.append("\n".join(rows))

    content = "\n\n".join(parts)
    if not content.strip():
        raise ValueError(
            "Document appears to be empty. No readable text found."
        )

    title = _title_from_text(content, fallback="Uploaded DOCX")
    return ExtractedDocument(
        title=title,
        content=content.strip(),
        file_type="docx",
        page_count=0,
        word_count=len(content.split()),
        char_count=len(content),
    )


def extract_markdown(text: str) -> ExtractedDocument:
    """Extract content from Markdown text, parsing YAML frontmatter."""
    if not text or not text.strip():
        raise ValueError("Empty markdown content")

    title = "Uploaded Markdown"
    content = text

    # Try parsing YAML frontmatter
    try:
        import frontmatter

        post = frontmatter.loads(text)
        content = post.content
        if post.get("title"):
            title = post["title"]
        elif post.get("name"):
            title = post["name"]
    except Exception:
        pass  # No frontmatter or parse error; use raw text

    # If no title from frontmatter, try first heading
    if title == "Uploaded Markdown":
        heading = re.search(r"^#\s+(.+)$", content, re.MULTILINE)
        if heading:
            title = heading.group(1).strip()

    return ExtractedDocument(
        title=title,
        content=content.strip(),
        file_type="md",
        page_count=0,
        word_count=len(content.split()),
        char_count=len(content),
    )


def _title_from_text(text: str, fallback: str = "Uploaded Document") -> str:
    """Extract a title from the first non-empty line of text."""
    for line in text.split("\n"):
        line = line.strip()
        if line and len(line) > 3:
            # Truncate long first lines
            return line[:120] if len(line) > 120 else line
    return fallback
