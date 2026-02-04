"""Document loader for DOCX and PDF files with chunking support.

Supports two modes:
1. Legacy mode: Simple character-based chunking (backward compatible)
2. Chunked mode: Uses hierarchical, structure-aware chunking (recommended)
"""

import logging
from dataclasses import dataclass, field
from pathlib import Path
from typing import Iterator, Optional

from .index_metadata_loader import get_document_metadata

# Import chunking module (optional, for enhanced chunking)
try:
    from ..chunking import (
        PolicyChunkingStrategy,
        ChunkingConfig,
        Chunk,
        ChunkedDocument,
    )
    CHUNKING_AVAILABLE = True
except ImportError:
    CHUNKING_AVAILABLE = False

logger = logging.getLogger(__name__)

# Maximum characters per chunk (roughly 1500 tokens = ~6000 chars)
MAX_CHUNK_SIZE = 6000
CHUNK_OVERLAP = 500


@dataclass
class PolicyDocument:
    """Represents a parsed policy document or chunk."""

    file_path: str
    title: str
    content: str
    doc_type: str = "policy"
    file_type: str = ""  # 'docx' or 'pdf'
    page_count: int = 0
    chunk_index: int = 0
    total_chunks: int = 1
    metadata: dict = field(default_factory=dict)

    # Ownership fields from index.md
    owner_team: str = ""
    owner_team_abbr: str = ""
    owner_department: str = ""
    owner_organization: str = ""
    owner_display: str = ""
    collection_name: str = ""

    def to_dict(self) -> dict:
        """Convert to dictionary for Weaviate ingestion."""
        # Add chunk info to title if multiple chunks
        title = self.title
        if self.total_chunks > 1:
            title = f"{self.title} (Part {self.chunk_index + 1}/{self.total_chunks})"

        full_text_parts = [f"Policy: {title}"]
        if self.owner_display:
            full_text_parts.append(f"Owner: {self.owner_display}")
        full_text_parts.append(f"\n{self.content}")

        return {
            "file_path": self.file_path,
            "title": title,
            "content": self.content,
            "file_type": self.file_type,
            "page_count": self.page_count,
            # Ownership fields
            "owner_team": self.owner_team,
            "owner_team_abbr": self.owner_team_abbr,
            "owner_department": self.owner_department,
            "owner_organization": self.owner_organization,
            "owner_display": self.owner_display,
            "collection_name": self.collection_name,
            "full_text": "\n".join(full_text_parts),
        }


def chunk_text(text: str, max_size: int = MAX_CHUNK_SIZE, overlap: int = CHUNK_OVERLAP) -> list[str]:
    """Split text into chunks with overlap.

    Args:
        text: Text to chunk
        max_size: Maximum characters per chunk
        overlap: Number of characters to overlap between chunks

    Returns:
        List of text chunks
    """
    if len(text) <= max_size:
        return [text]

    chunks = []
    start = 0

    while start < len(text):
        end = start + max_size

        # Try to break at a paragraph or sentence boundary
        if end < len(text):
            # Look for paragraph break
            para_break = text.rfind("\n\n", start, end)
            if para_break > start + max_size // 2:
                end = para_break + 2
            else:
                # Look for sentence break
                sentence_break = text.rfind(". ", start, end)
                if sentence_break > start + max_size // 2:
                    end = sentence_break + 2

        chunk = text[start:end].strip()
        if chunk:
            chunks.append(chunk)

        # Move start with overlap
        start = end - overlap if end < len(text) else len(text)

    return chunks


class DocumentLoader:
    """Loader for DOCX and PDF policy documents with chunking."""

    def __init__(self, documents_path: Path):
        """Initialize the document loader.

        Args:
            documents_path: Path to directory containing documents
        """
        self.documents_path = Path(documents_path)

    def load_all(self) -> Iterator[dict]:
        """Load all DOCX and PDF files, chunking large documents.

        Yields:
            Dictionary representations of parsed document chunks
        """
        # Load DOCX files
        docx_files = list(self.documents_path.glob("*.docx"))
        pdf_files = list(self.documents_path.glob("*.pdf"))

        logger.info(
            f"Found {len(docx_files)} DOCX and {len(pdf_files)} PDF files to process"
        )

        for docx_file in docx_files:
            try:
                for doc_chunk in self._load_docx_chunked(docx_file):
                    yield doc_chunk.to_dict()
            except Exception as e:
                logger.error(f"Error loading DOCX {docx_file}: {e}")
                continue

        for pdf_file in pdf_files:
            try:
                for doc_chunk in self._load_pdf_chunked(pdf_file):
                    yield doc_chunk.to_dict()
            except Exception as e:
                logger.error(f"Error loading PDF {pdf_file}: {e}")
                continue

    def _load_docx_chunked(self, file_path: Path) -> Iterator[PolicyDocument]:
        """Load a DOCX file and yield chunks.

        Args:
            file_path: Path to the DOCX file

        Yields:
            PolicyDocument chunks
        """
        try:
            from docx import Document
        except ImportError:
            logger.error("python-docx not installed. Run: pip install python-docx")
            return

        try:
            doc = Document(str(file_path))

            # Get ownership metadata from index.md
            index_metadata = get_document_metadata(file_path)

            # Extract text from paragraphs
            paragraphs = []
            for para in doc.paragraphs:
                text = para.text.strip()
                if text:
                    paragraphs.append(text)

            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = " | ".join(
                        cell.text.strip() for cell in row.cells if cell.text.strip()
                    )
                    if row_text:
                        paragraphs.append(row_text)

            content = "\n\n".join(paragraphs)
            title = self._extract_title(file_path, paragraphs)

            # Chunk the content
            chunks = chunk_text(content)
            total_chunks = len(chunks)

            for i, chunk_content in enumerate(chunks):
                yield PolicyDocument(
                    file_path=str(file_path),
                    title=title,
                    content=chunk_content,
                    file_type="docx",
                    chunk_index=i,
                    total_chunks=total_chunks,
                    owner_team=index_metadata.get("owner_team", ""),
                    owner_team_abbr=index_metadata.get("owner_team_abbr", ""),
                    owner_department=index_metadata.get("owner_department", ""),
                    owner_organization=index_metadata.get("owner_organization", ""),
                    owner_display=index_metadata.get("owner_display", ""),
                    collection_name=index_metadata.get("collection_name", ""),
                )

        except Exception as e:
            logger.error(f"Failed to parse DOCX {file_path}: {e}")

    def _load_pdf_chunked(self, file_path: Path) -> Iterator[PolicyDocument]:
        """Load a PDF file and yield chunks.

        Args:
            file_path: Path to the PDF file

        Yields:
            PolicyDocument chunks
        """
        # Try PyMuPDF first (better extraction)
        try:
            yield from self._load_pdf_pymupdf_chunked(file_path)
            return
        except ImportError:
            pass

        # Fall back to pypdf (successor to PyPDF2)
        try:
            yield from self._load_pdf_pypdf_chunked(file_path)
        except ImportError:
            logger.error(
                "Neither pymupdf nor pypdf installed. "
                "Run: pip install pymupdf or pip install pypdf"
            )

    def _load_pdf_pymupdf_chunked(self, file_path: Path) -> Iterator[PolicyDocument]:
        """Load PDF using PyMuPDF and yield chunks."""
        import fitz  # PyMuPDF

        try:
            doc = fitz.open(str(file_path))
            pages = []

            # Get ownership metadata from index.md
            index_metadata = get_document_metadata(file_path)

            for page in doc:
                text = page.get_text()
                if text.strip():
                    pages.append(text.strip())

            content = "\n\n".join(pages)
            title = self._extract_title(file_path, pages)
            page_count = len(doc)

            # Chunk the content
            chunks = chunk_text(content)
            total_chunks = len(chunks)

            for i, chunk_content in enumerate(chunks):
                yield PolicyDocument(
                    file_path=str(file_path),
                    title=title,
                    content=chunk_content,
                    file_type="pdf",
                    page_count=page_count,
                    chunk_index=i,
                    total_chunks=total_chunks,
                    owner_team=index_metadata.get("owner_team", ""),
                    owner_team_abbr=index_metadata.get("owner_team_abbr", ""),
                    owner_department=index_metadata.get("owner_department", ""),
                    owner_organization=index_metadata.get("owner_organization", ""),
                    owner_display=index_metadata.get("owner_display", ""),
                    collection_name=index_metadata.get("collection_name", ""),
                )

        except Exception as e:
            logger.error(f"Failed to parse PDF with PyMuPDF {file_path}: {e}")

    def _load_pdf_pypdf_chunked(self, file_path: Path) -> Iterator[PolicyDocument]:
        """Load PDF using pypdf and yield chunks."""
        from pypdf import PdfReader

        try:
            reader = PdfReader(str(file_path))
            pages = []

            # Get ownership metadata from index.md
            index_metadata = get_document_metadata(file_path)

            for page in reader.pages:
                text = page.extract_text()
                if text and text.strip():
                    pages.append(text.strip())

            content = "\n\n".join(pages)
            title = self._extract_title(file_path, pages)
            page_count = len(reader.pages)

            # Chunk the content
            chunks = chunk_text(content)
            total_chunks = len(chunks)

            for i, chunk_content in enumerate(chunks):
                yield PolicyDocument(
                    file_path=str(file_path),
                    title=title,
                    content=chunk_content,
                    file_type="pdf",
                    page_count=page_count,
                    chunk_index=i,
                    total_chunks=total_chunks,
                    owner_team=index_metadata.get("owner_team", ""),
                    owner_team_abbr=index_metadata.get("owner_team_abbr", ""),
                    owner_department=index_metadata.get("owner_department", ""),
                    owner_organization=index_metadata.get("owner_organization", ""),
                    owner_display=index_metadata.get("owner_display", ""),
                    collection_name=index_metadata.get("collection_name", ""),
                )

        except Exception as e:
            logger.error(f"Failed to parse PDF with pypdf {file_path}: {e}")

    def _extract_title(self, file_path: Path, paragraphs: list[str]) -> str:
        """Extract title from filename or first paragraph.

        Args:
            file_path: Path to the file
            paragraphs: List of extracted paragraphs

        Returns:
            Extracted title
        """
        # Try to get title from first short paragraph (likely a heading)
        for para in paragraphs[:3]:
            if len(para) < 100 and para.strip():
                return para.strip()

        # Fall back to filename
        return file_path.stem.replace("-", " ").replace("_", " ").title()

    # ========== Chunked Loading Methods (New) ==========

    def load_all_chunked(
        self,
        config: Optional["ChunkingConfig"] = None,
    ) -> Iterator["ChunkedDocument"]:
        """Load all documents with hierarchical, structure-aware chunking.

        This method creates multiple chunks per document:
        - Section-level chunks (detected from headers)
        - Granular semantic chunks for large sections

        Args:
            config: Optional chunking configuration

        Yields:
            ChunkedDocument objects with hierarchical chunks
        """
        if not CHUNKING_AVAILABLE:
            logger.warning("Chunking module not available. Use load_all() instead.")
            return

        strategy = PolicyChunkingStrategy(config)

        docx_files = list(self.documents_path.glob("*.docx"))
        pdf_files = list(self.documents_path.glob("*.pdf"))

        logger.info(
            f"Loading {len(docx_files)} DOCX and {len(pdf_files)} PDF files with chunking"
        )

        for docx_file in docx_files:
            try:
                content, title = self._extract_docx_content(docx_file)
                if content:
                    index_metadata = get_document_metadata(docx_file)
                    chunked_doc = strategy.chunk_document(
                        content=content,
                        file_path=str(docx_file),
                        title=title,
                        metadata=index_metadata,
                    )
                    logger.debug(
                        f"Chunked DOCX '{title}' into {chunked_doc.total_chunks} chunks"
                    )
                    yield chunked_doc
            except Exception as e:
                logger.error(f"Error chunking DOCX {docx_file}: {e}")
                continue

        for pdf_file in pdf_files:
            try:
                content, title = self._extract_pdf_content(pdf_file)
                if content:
                    index_metadata = get_document_metadata(pdf_file)
                    chunked_doc = strategy.chunk_document(
                        content=content,
                        file_path=str(pdf_file),
                        title=title,
                        metadata=index_metadata,
                    )
                    logger.debug(
                        f"Chunked PDF '{title}' into {chunked_doc.total_chunks} chunks"
                    )
                    yield chunked_doc
            except Exception as e:
                logger.error(f"Error chunking PDF {pdf_file}: {e}")
                continue

    def _extract_docx_content(self, file_path: Path) -> tuple[str, str]:
        """Extract content and title from a DOCX file.

        Args:
            file_path: Path to the DOCX file

        Returns:
            Tuple of (content, title)
        """
        try:
            from docx import Document
        except ImportError:
            logger.error("python-docx not installed. Run: pip install python-docx")
            return "", ""

        try:
            doc = Document(str(file_path))

            paragraphs = []
            for para in doc.paragraphs:
                text = para.text.strip()
                if text:
                    paragraphs.append(text)

            for table in doc.tables:
                for row in table.rows:
                    row_text = " | ".join(
                        cell.text.strip() for cell in row.cells if cell.text.strip()
                    )
                    if row_text:
                        paragraphs.append(row_text)

            content = "\n\n".join(paragraphs)
            title = self._extract_title(file_path, paragraphs)

            return content, title

        except Exception as e:
            logger.error(f"Failed to extract DOCX content {file_path}: {e}")
            return "", ""

    def _extract_pdf_content(self, file_path: Path) -> tuple[str, str]:
        """Extract content and title from a PDF file.

        Args:
            file_path: Path to the PDF file

        Returns:
            Tuple of (content, title)
        """
        # Try PyMuPDF first
        try:
            import fitz

            doc = fitz.open(str(file_path))
            pages = []

            for page in doc:
                text = page.get_text()
                if text.strip():
                    pages.append(text.strip())

            content = "\n\n".join(pages)
            title = self._extract_title(file_path, pages)
            return content, title

        except ImportError:
            pass
        except Exception as e:
            logger.error(f"PyMuPDF failed for {file_path}: {e}")

        # Fall back to pypdf
        try:
            from pypdf import PdfReader

            reader = PdfReader(str(file_path))
            pages = []

            for page in reader.pages:
                text = page.extract_text()
                if text and text.strip():
                    pages.append(text.strip())

            content = "\n\n".join(pages)
            title = self._extract_title(file_path, pages)
            return content, title

        except ImportError:
            logger.error("Neither pymupdf nor pypdf installed.")
            return "", ""
        except Exception as e:
            logger.error(f"pypdf failed for {file_path}: {e}")
            return "", ""
